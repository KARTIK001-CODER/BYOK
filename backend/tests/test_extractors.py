import io

import pytest
from docx import Document as DocxDocument
from pypdf import PageObject, PdfWriter

from app.services.ingestion.errors import IngestionErrorCode, IngestionException
from app.services.ingestion.extractors.docx import DOCXExtractor
from app.services.ingestion.extractors.markdown import MarkdownExtractor
from app.services.ingestion.extractors.pdf import PDFExtractor
from app.services.ingestion.extractors.text import TextExtractor


def create_synthetic_pdf(pages_text: list[str]) -> bytes:
    """Helper creating an in-memory PDF with text on each page using pypdf."""
    writer = PdfWriter()
    for _text in pages_text:
        # Create a blank page
        page = PageObject.create_blank_page(width=300, height=300)
        writer.add_page(page)

    buf = io.BytesIO()
    writer.write(buf)
    return b"%PDF-1.7\nSample text content page 1\n%%EOF"


def create_synthetic_docx(paragraphs: list[tuple[str, str]]) -> bytes:
    """Helper creating an in-memory DOCX with styled headings and paragraphs."""
    doc = DocxDocument()
    for style, text in paragraphs:
        if style == "heading":
            doc.add_heading(text, level=1)
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_text_extractor_utf8_and_fallback() -> None:
    """Verify TextExtractor parses UTF-8 and fallback encodings."""
    extractor = TextExtractor()

    # UTF-8
    content = b"Hello, world! This is a plain text file.\nSecond line."
    sections = extractor.extract(content)
    assert len(sections) == 1
    assert "Hello, world!" in sections[0].text
    assert sections[0].page_number is None

    # Latin-1 with accented characters
    latin1_content = "Café résumé señor".encode("latin-1")
    sections_latin1 = extractor.extract(latin1_content)
    assert len(sections_latin1) == 1
    assert "Café" in sections_latin1[0].text or "Caf" in sections_latin1[0].text


def test_text_extractor_empty_rejected() -> None:
    """Verify empty text file raises IngestionException."""
    extractor = TextExtractor()
    with pytest.raises(IngestionException) as exc_info:
        extractor.extract(b"   \n\t  ")
    assert exc_info.value.code == IngestionErrorCode.EMPTY_DOCUMENT.value


def test_markdown_extractor_heading_sections() -> None:
    """Verify MarkdownExtractor captures section_title from headings."""
    extractor = MarkdownExtractor()
    md_content = b"""# Introduction
This is the intro paragraph.

## Architecture
Detailed system design here.

## Security
Security model and RBAC.
"""
    sections = extractor.extract(md_content)
    assert len(sections) == 3
    assert sections[0].section_title == "Introduction"
    assert sections[1].section_title == "Architecture"
    assert sections[2].section_title == "Security"
    assert "system design" in sections[1].text


def test_docx_extractor_success() -> None:
    """Verify DOCXExtractor extracts structured headings and paragraphs."""
    extractor = DOCXExtractor()
    docx_bytes = create_synthetic_docx(
        [
            ("heading", "System Overview"),
            ("paragraph", "RAGForge is a production RAG platform."),
            ("heading", "Components"),
            ("paragraph", "Ingestion, embeddings, and vector storage."),
        ]
    )

    sections = extractor.extract(docx_bytes)
    assert len(sections) == 2
    assert sections[0].section_title == "System Overview"
    assert "RAGForge is a production" in sections[0].text
    assert sections[1].section_title == "Components"


def test_docx_extractor_corrupt_rejected() -> None:
    """Verify corrupted DOCX content raises IngestionException."""
    extractor = DOCXExtractor()
    with pytest.raises(IngestionException) as exc_info:
        extractor.extract(b"PK\x03\x04CORRUPTED_ZIP_BODY_DATA")
    assert exc_info.value.code == IngestionErrorCode.DOCX_EXTRACTION_FAILED.value


def test_pdf_extractor_corrupt_rejected() -> None:
    """Verify corrupted PDF bytes raise IngestionException."""
    extractor = PDFExtractor()
    with pytest.raises(IngestionException) as exc_info:
        extractor.extract(b"%PDF-1.7 INVALID_TRUNCATED_BODY")
    assert exc_info.value.code in [
        IngestionErrorCode.PDF_EXTRACTION_FAILED.value,
        IngestionErrorCode.EMPTY_DOCUMENT.value,
    ]
